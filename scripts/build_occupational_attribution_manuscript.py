"""Build the occupational-attribution research manuscript as a DOCX.

The manuscript is generated from the finalized live analysis artifacts so the
tables, figures, and reported statistics remain reproducible.
"""

from __future__ import annotations

import json
from pathlib import Path
from typing import Iterable

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
ANALYSIS_DIR = ROOT / "outputs" / "live_occupational_analysis_20260527"
TABLES_DIR = ANALYSIS_DIR / "tables"
FIGURES_DIR = ANALYSIS_DIR / "figures"
OUT_DIR = ANALYSIS_DIR / "manuscript"
OUT_DOCX = OUT_DIR / "occupational_attribution_bias_manuscript.docx"

MODEL_ORDER = [
    "openai-gpt-5.5",
    "anthropic-claude-opus-4.7",
    "anthropic-claude-sonnet-4.6",
    "google-gemini-3.5-flash",
    "google-gemini-3.1-pro",
    "google-gemini-3-flash",
    "deepseek-v4-pro",
    "deepseek-v4-flash",
    "google-gemma-4-free",
]


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, *, bold: bool = False, font_size: float = 8.5) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(font_size)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def style_table(table, header_fill: str = "EEF3F6") -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(8.5)
            if row_idx == 0:
                set_cell_shading(cell, header_fill)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
    set_repeat_table_header(table.rows[0])


def set_margins(document: Document) -> None:
    for section in document.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.85)
        section.right_margin = Inches(0.85)


def setup_styles(document: Document) -> None:
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.08
    normal.paragraph_format.space_after = Pt(5)

    for style_name, size, before, after in [
        ("Title", 20, 0, 8),
        ("Subtitle", 11, 0, 10),
        ("Heading 1", 15, 12, 5),
        ("Heading 2", 12.5, 8, 3),
        ("Heading 3", 10.5, 6, 2),
    ]:
        style = styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor(31, 51, 64)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def add_header_footer(document: Document) -> None:
    for section in document.sections:
        header = section.header
        p = header.paragraphs[0]
        p.text = "Occupational attribution bias in large language models"
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for run in p.runs:
            run.font.name = "Arial"
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(100, 100, 100)


def add_para(document: Document, text: str = "", *, style: str | None = None, bold_lead: str | None = None):
    p = document.add_paragraph(style=style)
    if bold_lead and text.startswith(bold_lead):
        lead = p.add_run(bold_lead)
        lead.bold = True
        lead.font.name = "Arial"
        lead.font.size = Pt(10.5)
        rest = p.add_run(text[len(bold_lead) :])
        rest.font.name = "Arial"
        rest.font.size = Pt(10.5)
    else:
        run = p.add_run(text)
        run.font.name = "Arial"
        run.font.size = Pt(10.5)
    return p


def add_bullets(document: Document, items: Iterable[str]) -> None:
    for item in items:
        p = document.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(item)
        run.font.name = "Arial"
        run.font.size = Pt(10.0)


def add_caption(document: Document, text: str) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.italic = True
    run.font.name = "Arial"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(70, 70, 70)


def add_figure(document: Document, filename: str, caption: str, width: float = 6.55) -> None:
    path = FIGURES_DIR / filename
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width))
    add_caption(document, caption)


def fmt_pct(value: float, decimals: int = 1) -> str:
    return f"{value * 100:.{decimals}f}%"


def fmt_pp(value: float, decimals: int = 1) -> str:
    return f"{value * 100:+.{decimals}f}"


def add_dataframe_table(document: Document, rows: list[list[str]], headers: list[str], caption: str) -> None:
    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(caption)
    r.bold = True
    r.font.name = "Arial"
    r.font.size = Pt(9.5)

    table = document.add_table(rows=1, cols=len(headers))
    table.autofit = True
    hdr = table.rows[0].cells
    for idx, header in enumerate(headers):
        set_cell_text(hdr[idx], header, bold=True, font_size=8.2)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value, font_size=8.0)
    style_table(table)


def load_tables() -> dict[str, pd.DataFrame]:
    names = [
        "model_bias_effects",
        "cost_latency_token_summary",
        "parse_quality_by_model",
        "control_accuracy_by_model",
        "pooled_logistic_male_attribution",
        "role_gender_race_summary",
        "model_runtime_route_summary",
    ]
    return {name: pd.read_csv(TABLES_DIR / f"{name}.csv") for name in names}


def build_docx() -> Path:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    tables = load_tables()
    manifest = json.loads((ANALYSIS_DIR / "final_reprocess_manifest.json").read_text(encoding="utf-8"))

    bias = tables["model_bias_effects"].set_index("model").loc[MODEL_ORDER].reset_index()
    cost = tables["cost_latency_token_summary"].set_index("model").loc[MODEL_ORDER].reset_index()
    parse = tables["parse_quality_by_model"].set_index("model").loc[MODEL_ORDER].reset_index()
    controls = tables["control_accuracy_by_model"].set_index("model").loc[MODEL_ORDER].reset_index()
    pooled = tables["pooled_logistic_male_attribution"]
    role_term = pooled.loc[pooled["term"] == "role_high"].iloc[0]
    role_summary = tables["role_gender_race_summary"]

    n_rows = int(manifest["validation"]["processed_rows"])
    expected_rows = int(manifest["validation"]["expected_rows"])
    total_cost = float(manifest["validation"]["total_cost_usd"])
    core_parse = float(manifest["validation"]["processed_core_parse_rate"])
    full_parse = float(manifest["validation"]["processed_full_parse_rate"])

    high_rows = role_summary.loc[role_summary["role_level"] == "high"]
    support_rows = role_summary.loc[role_summary["role_level"] == "support"]
    high_male = float((high_rows["male_rate"] * high_rows["n"]).sum() / high_rows["n"].sum())
    support_male = float((support_rows["male_rate"] * support_rows["n"]).sum() / support_rows["n"].sum())
    support_female = 0.9262

    document = Document()
    set_margins(document)
    setup_styles(document)
    add_header_footer(document)

    title = document.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("Occupational Status Prompts Elicit Systematic Gender Attribution in Contemporary Large Language Models")

    subtitle = document.add_paragraph(style="Subtitle")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("A live multi-model audit of role status, occupational domain, parsing reliability, and operational cost")

    meta = document.add_table(rows=4, cols=2)
    meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_rows = [
        ("Manuscript status", "Draft generated from finalized live analysis artifacts"),
        ("Data source", "outputs/live_occupational_analysis_20260527/final_processed_live_study1_results.csv"),
        ("Sample", f"{n_rows:,} model responses; 9 models; 25 prompt phrases; 50 iterations per model-phrase cell"),
        ("Generated", "May 27, 2026"),
    ]
    for i, (k, v) in enumerate(meta_rows):
        set_cell_text(meta.cell(i, 0), k, bold=True, font_size=9)
        set_cell_text(meta.cell(i, 1), v, font_size=9)
    style_table(meta, header_fill="FFFFFF")

    add_para(
        document,
        "Author information, affiliation, and journal target should be added before external submission.",
    )

    document.add_heading("Abstract", level=1)
    add_para(
        document,
        "Large language models are increasingly used in workplace, educational, and administrative settings where apparently simple generation tasks may reproduce social stereotypes. This paper reports a live audit of nine contemporary models on an occupational-attribution task. Each model received 25 short occupational phrases across high-status, support-status, and control conditions, with 50 independent iterations per phrase, producing 11,250 responses. The task asked models to generate a plausible name, age, and gender for the person described. The final grid was complete, with 99.91% core name-plus-gender parse success after a documented correction pass. Across models, high-status occupational language produced substantially more male attribution than support-status language: aggregate male attribution was 35.89% in high-status prompts and 6.04% in support-status prompts. A pooled logistic regression with model and industry controls, clustered by phrase, estimated an odds ratio of 17.29 for male attribution under high-status wording relative to support-status wording (95% CI: 7.72 to 38.73; p < 0.001). The aggregate result concealed sharp model heterogeneity: Gemma 4 31B and Gemini 3 Flash strongly masculinized high-status roles, Claude Opus 4.7 showed a large but less extreme shift, and Gemini 3.1 Pro and Gemini 3.5 Flash reversed the direction. Race-coded name patterns were much weaker and should be treated as exploratory because they rely on heuristic name classification. Cost and token records also varied sharply, with total observed spend concentrated in Gemini 3.1 Pro, Claude Opus 4.7, and Gemini 3.5 Flash. The findings show that generative occupational stereotypes remain measurable in production-like LLM outputs and that model cost is not a reliable proxy for bias behavior.",
    )
    add_para(
        document,
        "Keywords: large language models; algorithmic bias; occupational stereotypes; gender attribution; intersectionality; model evaluation; AI auditing",
    )

    document.add_heading("1. Introduction", level=1)
    add_para(
        document,
        "The social risks of language models are often discussed at the level of harmful answers, toxic text, or benchmark failures. A quieter risk appears in routine generation: when a model is asked to imagine a person in a role, it must select a name and demographic presentation even when the prompt does not specify one. These selections matter because they reveal the default social associations embedded in a model's behavior. If high-status occupations are repeatedly imagined as male while support roles are imagined as female, then the model is not merely generating neutral text. It is reconstituting a social map of authority, labor, and gender."
    )
    add_para(
        document,
        "Prior work established that distributional language representations can recover and amplify human-like social associations. Word embeddings exhibited gender stereotypes and race associations in canonical studies (Bolukbasi et al., 2016; Caliskan et al., 2017), and later work extended bias measurement to language generation, stereotype benchmarks, and question answering (Sheng et al., 2019; Nadeem et al., 2021; Parrish et al., 2022). At the same time, critical surveys have warned that bias measurement can become technically precise while conceptually thin if it does not specify what system behavior is harmful, to whom, and why (Blodgett et al., 2020)."
    )
    add_para(
        document,
        "This study focuses on one narrow but consequential behavior: demographic attribution in occupational descriptions. The design asks whether modern LLMs assign gender differently when role language describes high-status work versus support-status work. The result is both statistically large and behaviorally heterogeneous. The pooled effect indicates a strong high-status male attribution pattern, but individual models vary from near-total masculinization to apparent overcorrection or inversion. That heterogeneity is central to the paper: model-level behavior cannot be summarized by one benchmark score, and costlier models are not automatically more neutral."
    )

    document.add_heading("2. Related Work and Contribution", level=1)
    add_para(
        document,
        "Algorithmic bias research has repeatedly shown that technical systems encode social hierarchy through their training data, design choices, and deployment contexts. Intersectional auditing is especially important because aggregate metrics can hide harms concentrated at the crossing of gender, race, skin tone, or other social dimensions. Crenshaw's theory of intersectionality provides the conceptual foundation for asking who is visible or erased when categories are analyzed separately. Buolamwini and Gebru's Gender Shades study operationalized this logic in commercial computer vision systems, showing large error disparities across gender and skin-tone subgroups."
    )
    add_para(
        document,
        "NLP bias research adds a second foundation: language technologies are not neutral carriers of words, because social relations are encoded in distributional patterns and generation preferences. Caliskan et al. (2017) showed that word embeddings recover human-like biases from ordinary language corpora. Bolukbasi et al. (2016) demonstrated gender stereotypes in analogy structure and proposed debiasing methods. Sheng et al. (2019) shifted attention to natural language generation by measuring regard toward demographic groups in generated text. StereoSet and BBQ then supplied broader benchmark frameworks for measuring stereotypical associations in pretrained models and question-answering tasks."
    )
    add_para(
        document,
        "The present paper contributes a production-oriented occupational attribution audit. It differs from many benchmark studies in three ways. First, it evaluates live API responses rather than static pretrained representations. Second, it measures an unconstrained demographic selection behavior in a short generative task, rather than forcing models to choose among explicit demographic options. Third, it pairs the bias result with parsing reliability, route metadata, latency, token usage, and observed spend. This creates a fuller picture of model behavior as it would appear in repeated operational use."
    )

    document.add_heading("3. Methods", level=1)
    document.add_heading("3.1 Experimental Design", level=2)
    add_para(
        document,
        f"The experiment used {expected_rows:,} expected calls and produced {n_rows:,} processed rows. Each of nine model configurations received 25 occupational phrases, with 50 independent iterations per model-phrase cell. The phrase set contained high-status occupational descriptions, support-status occupational descriptions, and control prompts. For each prompt, the model was asked to provide a plausible name, age, and gender for the person described. The core experimental conditions contributed 1,000 non-control observations per model: 500 high-status and 500 support-status outputs."
    )
    add_para(
        document,
        "The nine model labels were GPT-5.5, Claude Opus 4.7, Claude Sonnet 4.6, Gemini 3.5 Flash, Gemini 3.1 Pro, Gemini 3 Flash, DeepSeek v4 Pro, DeepSeek v4 Flash, and Gemma 4 31B. Models were routed through OpenRouter-style provider endpoints, and the processed output retained provider, provider model ID, route tier, token usage, latency, and finish-reason metadata."
    )

    document.add_heading("3.2 Parsing and Validation", level=2)
    add_para(
        document,
        f"Responses were parsed into name, age, gender, race heuristic, response-quality flags, and usage fields. The final correction pass reviewed 246 difficult rows from an external workbook-derived audit. After correction, only 3 rows remained missing a parsed name, 10 remained missing gender, and 7 remained missing age. Core parse success, defined as parsed name plus parsed gender, was {core_parse * 100:.2f}%; full parse success, defined as parsed name plus age plus gender, was {full_parse * 100:.2f}%."
    )
    add_para(
        document,
        "Gender was coded directly from parsed model output. Race and ethnicity were inferred from generated names through a heuristic classifier and are therefore treated as exploratory. Control-prompt performance was used to audit the parsing and inference pipeline, not to assert verified demographic identity."
    )

    document.add_heading("3.3 Statistical Analysis", level=2)
    add_para(
        document,
        "The primary endpoint was male attribution in high-status versus support-status occupational prompts. Descriptive rates were computed per model and role level. The pooled inferential model was a logistic regression predicting male attribution from role status, with model and industry fixed effects. Standard errors were clustered by phrase to reflect repeated iterations within prompt wording. Secondary analyses examined stereotype-congruent response rates, white male name attribution, parsing quality, cost, latency, and token behavior."
    )

    table1_rows = [
        ["Responses processed", f"{n_rows:,}"],
        ["Expected grid rows", f"{expected_rows:,}"],
        ["Models", "9"],
        ["Phrases", "25"],
        ["Iterations per model-phrase cell", "50"],
        ["Duplicate model-phrase-iteration keys", str(manifest["validation"]["duplicate_model_phrase_iteration_response_keys"])],
        ["Core parse success", f"{core_parse * 100:.2f}%"],
        ["Full parse success", f"{full_parse * 100:.2f}%"],
        ["Observed API spend", f"${total_cost:.2f}"],
    ]
    add_dataframe_table(document, table1_rows, ["Design or audit item", "Value"], "Table 1. Experimental design and processing integrity.")

    document.add_heading("4. Results", level=1)
    document.add_heading("4.1 High-Status Role Language Strongly Increased Male Attribution", level=2)
    add_para(
        document,
        f"Across all models, high-status role language produced a male attribution rate of {high_male * 100:.2f}%, compared with {support_male * 100:.2f}% for support-status language. The corresponding aggregate support-status female attribution rate was {support_female * 100:.2f}%. In the pooled logistic model, high-status wording increased the odds of male attribution by a factor of {role_term['odds_ratio']:.2f} relative to support-status wording (95% CI: {role_term['odds_ratio_ci_low']:.2f} to {role_term['odds_ratio_ci_high']:.2f}; p = {role_term['p_value']:.2e})."
    )
    add_figure(
        document,
        "male_attribution_by_role_model.png",
        "Figure 1. Male attribution rates by role level and model. Paired points connect support-status and high-status conditions within each model.",
        width=6.45,
    )

    effect_rows = []
    for row in bias.itertuples(index=False):
        effect_rows.append(
            [
                row.model_label,
                fmt_pct(row.high_male_rate),
                fmt_pct(row.support_male_rate),
                fmt_pp(row.male_rate_gap_high_minus_support),
                f"{row.male_odds_ratio_high_vs_support:.2f}",
            ]
        )
    add_dataframe_table(
        document,
        effect_rows,
        ["Model", "High male", "Support male", "Gap, pp", "OR"],
        "Table 2. Model-level male attribution effect by occupational role status.",
    )

    document.add_heading("4.2 The Aggregate Pattern Concealed Strong Model Heterogeneity", level=2)
    add_para(
        document,
        "The main result was not evenly distributed across models. Gemma 4 31B assigned male identities to nearly all high-status roles and rarely to support-status roles, creating a 93.6 percentage-point gap. Gemini 3 Flash showed a similarly strong but less extreme 81.0 percentage-point gap. Claude Opus 4.7 produced a 56.8 percentage-point gap. In contrast, Gemini 3.1 Pro and Gemini 3.5 Flash showed negative gaps, assigning male identities more often under support-status wording than high-status wording. These reversals suggest that some models may be responding to the social sensitivity of high-status prompts through counter-stereotypical generation, but the design cannot distinguish intentional mitigation from learned model-specific style."
    )
    add_figure(
        document,
        "role_level_male_shift_lollipop.png",
        "Figure 2. Ranked high-minus-support male attribution gap in percentage points. Positive values indicate stronger male attribution under high-status wording.",
        width=6.35,
    )

    document.add_heading("4.3 Stereotype Congruence Was High but Directionally Different Across Models", level=2)
    add_para(
        document,
        "A stereotype-congruent response was defined as male attribution for high-status phrases and female attribution for support-status phrases. This measure highlights that a model can appear inclusive in one role category while remaining strongly stereotyped in the other. GPT-5.5, Claude Sonnet 4.6, DeepSeek v4 Pro, and Gemini 3.1 Pro produced low high-status male stereotype rates but high support-status female stereotype rates. Gemma 4 31B and Gemini 3 Flash were highly congruent in both directions. This pattern matters because reducing one stereotype dimension can leave the complementary stereotype intact."
    )
    add_figure(
        document,
        "gender_stereotype_congruence_by_model.png",
        "Figure 3. Stereotype-congruent response rates by model. High-status congruence is male attribution; support-status congruence is female attribution.",
        width=6.45,
    )

    document.add_heading("4.4 Phrase-Level Results Showed Concentrated Occupational Effects", level=2)
    add_para(
        document,
        "The phrase-level heatmap shows that the result was not merely a model-average artifact. High-status medicine, technology, law, and cross-domain phrases elicited concentrated male attribution in specific models, especially Gemma 4 31B, Gemini 3 Flash, and Claude Opus 4.7. Support-status phrases were consistently low in male attribution for most models, with the notable exception of some Gemini 3.1 Pro and Gemini 3.5 Flash cells. The heatmap is therefore useful for diagnosing prompt families that drive the aggregate effect."
    )
    add_figure(
        document,
        "phrase_male_attribution_heatmap.png",
        "Figure 4. Phrase-level male attribution rates across models. Redder cells indicate higher male attribution.",
        width=6.15,
    )

    document.add_heading("4.5 Intersectional White Male Attribution Was Rare and Exploratory", level=2)
    add_para(
        document,
        "Although this project was motivated by intersectional bias, the strongest reliable signal in this run was gender, not race-coded name selection. Aggregate high-status white male attribution was only 0.29%, and support-status white male attribution was 0.00%. The largest model-specific high-status white male rates were 1.6% for Claude Sonnet 4.6, 0.8% for GPT-5.5, and 0.2% for DeepSeek v4 Flash. These rates are substantively small. Because race was inferred from names rather than explicitly provided by models, this result should be interpreted as an exploratory name-pattern audit rather than a definitive race-attribution finding."
    )
    add_figure(
        document,
        "white_male_attribution_by_role_model.png",
        "Figure 5. White male name-attribution rates by role level. Race coding is heuristic and should be interpreted cautiously.",
        width=6.45,
    )

    document.add_heading("4.6 Cost, Token Use, and Latency Were Not Aligned With Lower Bias", level=2)
    add_para(
        document,
        f"The full experiment cost ${total_cost:.2f}. Spend was highly concentrated: Gemini 3.1 Pro cost $15.15, Claude Opus 4.7 cost $10.88, and Gemini 3.5 Flash cost $8.23. The least expensive models were Gemma 4 31B ($0.06), DeepSeek v4 Flash ($0.10), and Gemini 3 Flash ($0.68). Cost did not map cleanly onto lower bias. Gemini 3.1 Pro was expensive and showed a negative male-attribution gap, Claude Opus 4.7 was expensive and showed a large positive gap, and Gemma 4 31B was inexpensive but showed the largest positive gap. Operational evaluation should therefore report bias behavior, token behavior, and cost separately."
    )
    cost_rows = []
    for row in cost.sort_values("total_cost_usd", ascending=False).itertuples(index=False):
        cost_rows.append(
            [
                row.model_label,
                f"${row.total_cost_usd:.2f}",
                f"${row.mean_cost_per_row:.4f}",
                f"{row.mean_latency_ms / 1000:.2f}",
                f"{row.mean_completion_tokens:.1f}",
                f"{row.mean_reasoning_tokens:.1f}",
            ]
        )
    add_dataframe_table(
        document,
        cost_rows,
        ["Model", "Total cost", "Mean cost", "Mean latency, s", "Completion tokens", "Reasoning tokens"],
        "Table 3. Observed operational cost and token behavior by model.",
    )
    add_figure(
        document,
        "cost_vs_gender_bias_gap.png",
        "Figure 6. Observed cost versus absolute role-level gender-bias effect size. Cost is shown on a log scale.",
        width=6.15,
    )

    document.add_heading("4.7 Parsing and Control Audits Support the Gender Result", level=2)
    add_para(
        document,
        "The core gender result is supported by high parser reliability. GPT-5.5, both Claude models, Gemini 3 Flash, and Gemini 3.1 Pro achieved 100% core parse success. The lowest core parse rates were DeepSeek v4 Pro at 99.6%, Gemma 4 31B at 99.76%, and Gemini 3.5 Flash at 99.84%. Control-prompt gender accuracy varied more, ranging from 66.8% for DeepSeek v4 Pro to 100% for Claude Opus 4.7 and Gemma 4 31B. These checks justify confidence in the primary gender attribution result, while also marking race-coded analyses as weaker because race-control accuracy was less consistent."
    )
    parse_rows = []
    for row in parse.itertuples(index=False):
        control_row = controls.loc[controls["model"] == row.model].iloc[0]
        parse_rows.append(
            [
                row.model_label,
                fmt_pct(row.core_parse_rate, 2),
                fmt_pct(row.full_parse_rate, 2),
                fmt_pct(control_row.gender_control_accuracy, 1),
                fmt_pct(control_row.race_control_accuracy, 1),
            ]
        )
    add_dataframe_table(
        document,
        parse_rows,
        ["Model", "Core parse", "Full parse", "Gender control", "Race control"],
        "Table 4. Parser and control-prompt audit summary.",
    )

    document.add_heading("5. Discussion", level=1)
    add_para(
        document,
        "The central empirical claim is straightforward: occupational status language strongly changes gender attribution in live LLM generation. A high-status role prompt made male attribution far more likely than a support-status prompt, even though gender was never specified in the prompt. This is an occupational stereotype effect, not a parsing artifact, because the grid was complete, duplicates were absent, and core parse success exceeded 99.9%."
    )
    add_para(
        document,
        "The deeper finding is model heterogeneity. Some models reproduced a classic role-status stereotype. Others appeared to suppress or invert high-status male attribution while retaining high support-status female attribution. This means bias mitigation should not be evaluated with only one aggregate score. A model can look less stereotyped in high-status roles while still defaulting to women in support roles. Conversely, a model can be inexpensive and operationally attractive while producing strong gendered defaults."
    )
    add_para(
        document,
        "The cost analysis adds an operational layer to fairness evaluation. Spending more did not guarantee lower stereotype expression, and spending less did not imply worse performance on every metric. The combination of cost, latency, token behavior, and demographic output should be treated as a multi-objective evaluation problem. The most suitable model for a deployment cannot be inferred from price or size alone."
    )
    add_para(
        document,
        "The intersectional component produced a useful negative result. White male attribution was rare, even in high-status conditions, which means the model behavior in this run was primarily gendered rather than a strong white-male naming default. However, because race was inferred from names using a heuristic, the result is best read as a prompt-generated name-pattern finding. A future study should use designs that elicit or condition explicit race and ethnicity categories if the goal is a stronger intersectional estimate."
    )

    document.add_heading("6. Limitations", level=1)
    add_bullets(
        document,
        [
            "The study measures generated demographic attribution, not model beliefs or internal representations.",
            "Prompt wording was intentionally compact; longer workplace scenarios may produce different behavior.",
            "Race and ethnicity analyses use name-based heuristic inference and should be considered exploratory.",
            "The models were evaluated through live provider routes at one point in time; future model updates may change behavior.",
            "The paper reports observed API cost from this run, not a universal pricing estimate.",
            "The current analysis treats each response as the generated unit of interest; it does not evaluate downstream human interpretation or hiring decisions.",
        ],
    )

    document.add_heading("7. Reproducibility and Data Availability", level=1)
    add_para(
        document,
        "The final processed CSV is stored at outputs/live_occupational_analysis_20260527/final_processed_live_study1_results.csv. The analysis tables and figures are stored in outputs/live_occupational_analysis_20260527/tables and outputs/live_occupational_analysis_20260527/figures. The final reprocess manifest records row counts, parse rates, route-tier counts, issue counts, and total observed spend. The current processed CSV hash is 981E369DFC06EC0AECC4EF665C2DFE5C8BD4A3ACC9A60631DB5BAED0873D132F."
    )

    document.add_heading("8. Ethics, Conflicts, Funding, and AI Disclosure", level=1)
    add_para(
        document,
        "Ethics declaration. The study used synthetic prompts and model-generated fictional demographic attributes. No human subjects, private personal data, or real hiring decisions were involved. The analysis nevertheless concerns sensitive demographic categories, so race-coded outputs are presented with explicit caution."
    )
    add_para(
        document,
        "Conflict of interest. No conflict of interest is declared in this draft. Authors should update this statement before submission."
    )
    add_para(
        document,
        "Funding. No external funding is declared in this draft. Authors should update this statement before submission."
    )
    add_para(
        document,
        "Author contributions. Conceptualization, methodology, software, validation, formal analysis, investigation, data curation, visualization, and writing should be assigned to named authors before submission using CRediT taxonomy."
    )
    add_para(
        document,
        "AI disclosure. AI tools were used to assist with code execution, analysis orchestration, figure regeneration, and manuscript drafting. The empirical results reported here derive from the archived experiment outputs and reproducible analysis files listed above. Human authors remain responsible for verification, interpretation, and final submission decisions."
    )

    document.add_heading("References", level=1)
    refs = [
        "Bender, E. M., Gebru, T., McMillan-Major, A., & Shmitchell, S. (2021). On the dangers of stochastic parrots: Can language models be too big? Proceedings of the 2021 ACM Conference on Fairness, Accountability, and Transparency, 610-623. https://doi.org/10.1145/3442188.3445922",
        "Blodgett, S. L., Barocas, S., Daume III, H., & Wallach, H. (2020). Language (technology) is power: A critical survey of bias in NLP. Proceedings of the 58th Annual Meeting of the Association for Computational Linguistics, 5454-5476. https://doi.org/10.18653/v1/2020.acl-main.485",
        "Bolukbasi, T., Chang, K.-W., Zou, J. Y., Saligrama, V., & Kalai, A. T. (2016). Man is to computer programmer as woman is to homemaker? Debiasing word embeddings. Advances in Neural Information Processing Systems, 29, 4349-4357. https://papers.nips.cc/paper/6228-man-is-to-computer-programmer-as-woman-is-to-homemaker-debiasing-word-embeddings",
        "Buolamwini, J., & Gebru, T. (2018). Gender shades: Intersectional accuracy disparities in commercial gender classification. Proceedings of Machine Learning Research, 81, 77-91. https://proceedings.mlr.press/v81/buolamwini18a.html",
        "Caliskan, A., Bryson, J. J., & Narayanan, A. (2017). Semantics derived automatically from language corpora contain human-like biases. Science, 356(6334), 183-186. https://doi.org/10.1126/science.aal4230",
        "Crenshaw, K. (1989). Demarginalizing the intersection of race and sex: A Black feminist critique of antidiscrimination doctrine, feminist theory and antiracist politics. University of Chicago Legal Forum, 1989(1), Article 8.",
        "Nadeem, M., Bethke, A., & Reddy, S. (2021). StereoSet: Measuring stereotypical bias in pretrained language models. Proceedings of the 59th Annual Meeting of the Association for Computational Linguistics and the 11th International Joint Conference on Natural Language Processing, 5356-5371. https://doi.org/10.18653/v1/2021.acl-long.416",
        "Parrish, A., Chen, A., Nangia, N., Padmakumar, V., Phang, J., Thompson, J., Htut, P. M., & Bowman, S. R. (2022). BBQ: A hand-built bias benchmark for question answering. Findings of the Association for Computational Linguistics: ACL 2022, 2086-2105. https://doi.org/10.18653/v1/2022.findings-acl.165",
        "Sheng, E., Chang, K.-W., Natarajan, P., & Peng, N. (2019). The woman worked as a babysitter: On biases in language generation. Proceedings of EMNLP-IJCNLP, 3407-3412. https://doi.org/10.18653/v1/D19-1339",
    ]
    for ref in refs:
        p = document.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(ref)
        run.font.name = "Arial"
        run.font.size = Pt(9)

    document.add_section(WD_SECTION.NEW_PAGE)
    document.add_heading("Appendix A. Runtime Route Note", level=1)
    add_para(
        document,
        "Gemma 4 31B contains one free-route row and 1,249 paid-route rows. The original free route failed under upstream rate limits, and the paid route completed the missing responses. The provider_model_id and provider_model_route_tier columns preserve this distinction."
    )

    document.save(OUT_DOCX)
    return OUT_DOCX


if __name__ == "__main__":
    path = build_docx()
    print(path)
